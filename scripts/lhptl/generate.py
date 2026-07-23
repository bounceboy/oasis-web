#!/usr/bin/env python3
"""
LHPTL Word generator — template-based, rule-based approach.
Baca JSON dari stdin, fill template DOCX, output ke stdout.

Usage:
  python3 generate.py < data.json > output.docx
  python3 generate.py /path/to/data.json  (output to stdout)
"""

import sys
import json
import copy
import io
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

TEMPLATE = Path(__file__).parent.parent.parent / 'templates' / 'lhptl-template.docx'
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NA = '[HARAP DIISI PENGAWAS]'


# ─── Helpers ──────────────────────────────────────────────────────────────────

def fmt_rupiah(n):
    if n is None:
        return '-'
    try:
        return f"{int(n):,}".replace(',', '.')
    except Exception:
        return str(n)


def fmt_pct(n, decimals=2):
    if n is None:
        return '-'
    try:
        return f"{float(n):.{decimals}f}%"
    except Exception:
        return str(n)


def terbilang_total(n_per, n_bh):
    """Format jumlah rekanan: total (X perorangan dan Y badan hukum)."""
    total = (n_per or 0) + (n_bh or 0)
    if total == 0:
        return NA
    return (
        f"{total}, terdiri dari {n_per or 0} perorangan"
        f" dan {n_bh or 0} badan hukum"
    )


def para_full_text(para):
    """Return full text of paragraph including tabs from <w:tab/> elements."""
    parts = []
    for r in para._element.findall(f'{{{W}}}r'):
        for child in r:
            if child.tag == f'{{{W}}}t':
                parts.append(child.text or '')
            elif child.tag == f'{{{W}}}tab':
                parts.append('\t')
    return ''.join(parts)


def set_last_wt_value(para, new_value):
    """
    Replace the text of the last <w:t> element in the paragraph.
    Used for KV paragraphs where label+TAB+': Value' is in run[0].
    """
    runs = para._element.findall(f'{{{W}}}r')
    last_t = None
    for r in runs:
        for t in r.findall(f'{{{W}}}t'):
            last_t = t
    if last_t is not None:
        # The last w:t in a KV para contains ': OldValue'
        existing = last_t.text or ''
        colon_idx = existing.find(': ')
        if colon_idx != -1:
            prefix = existing[: colon_idx + 2]  # keep ': '
        else:
            colon_idx = existing.find(':')
            prefix = existing[: colon_idx + 1] + ' ' if colon_idx != -1 else ''
        last_t.text = prefix + new_value
        if new_value and (new_value[0] == ' ' or new_value[-1] == ' '):
            last_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def replace_full_para(para, new_text):
    """
    Replace entire paragraph text with new_text.
    Keeps formatting of the first run; discards subsequent runs.
    """
    runs = para._element.findall(f'{{{W}}}r')
    if not runs:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = new_text
        r.append(t)
        para._element.append(r)
        return

    # Keep first run, replace its text entirely
    r0 = runs[0]
    # Remove all w:t and w:tab children from first run
    for child in list(r0):
        if child.tag in (f'{{{W}}}t', f'{{{W}}}tab'):
            r0.remove(child)
    t = OxmlElement('w:t')
    t.text = new_text
    if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r0.append(t)

    # Remove remaining runs
    for r in runs[1:]:
        para._element.remove(r)


def find_para_by_start(paragraphs, prefix_lower):
    """Return first paragraph whose text starts with prefix_lower (case-insensitive)."""
    for p in paragraphs:
        if para_full_text(p).lower().lstrip().startswith(prefix_lower):
            return p
    return None


def find_para_exact(paragraphs, text_lower):
    """Return first paragraph whose stripped text equals text_lower (case-insensitive)."""
    for p in paragraphs:
        if para_full_text(p).strip().lower() == text_lower:
            return p
    return None


# ─── Table helpers ────────────────────────────────────────────────────────────

def get_cell_rPr(cell):
    """Get rPr from first run in first paragraph of a table cell (for font cloning)."""
    for p in cell.paragraphs:
        for r in p.runs:
            return r._r.find(f'{{{W}}}rPr')
    return None


def make_cell_para(text, rPr_src=None, bold=False, center=False):
    """Create a paragraph element for a table cell."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(f'{{{W}}}val', 'TableParagraph')
    pPr.append(pStyle)
    if center:
        jc = OxmlElement('w:jc')
        jc.set(f'{{{W}}}val', 'center')
        pPr.append(jc)
    p.append(pPr)

    r = OxmlElement('w:r')
    if rPr_src is not None:
        rPr = copy.deepcopy(rPr_src)
        # Set/remove bold
        b_el = rPr.find(f'{{{W}}}b')
        if bold and b_el is None:
            b_new = OxmlElement('w:b')
            rPr.insert(0, b_new)
        elif not bold and b_el is not None:
            rPr.remove(b_el)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text if text else '-'
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


def clone_row_format(source_row):
    """Clone a table row element (keeping cell formatting but clearing text)."""
    new_tr = copy.deepcopy(source_row._tr)
    # Clear text in all cells
    for tc in new_tr.findall(f'{{{W}}}tc'):
        for p in tc.findall(f'{{{W}}}p'):
            for r in p.findall(f'{{{W}}}r'):
                for t in r.findall(f'{{{W}}}t'):
                    t.text = ''
    return new_tr


def add_data_row(table, values, header_row_idx=0):
    """
    Add a new data row to table by cloning the header row format,
    then filling each cell with values[i].
    """
    src_row = table.rows[header_row_idx]
    rPr_list = [get_cell_rPr(cell) for cell in src_row.cells]

    new_tr = clone_row_format(src_row)
    cells = new_tr.findall(f'{{{W}}}tc')

    for i, cell_el in enumerate(cells):
        val = values[i] if i < len(values) else ''
        # Clear all paragraphs in cell
        for p in cell_el.findall(f'{{{W}}}p'):
            cell_el.remove(p)
        rPr = rPr_list[i] if i < len(rPr_list) else None
        cell_el.append(make_cell_para(val, rPr_src=rPr))

    table._tbl.append(new_tr)


def clear_data_rows(table, keep=1):
    """Remove all rows beyond the first `keep` rows."""
    rows = table.rows
    for row in rows[keep:]:
        table._tbl.remove(row._tr)


def update_header_cell(table, row_idx, col_idx, new_text):
    """Update a header cell text (e.g., year in lapkeu table)."""
    cell = table.rows[row_idx].cells[col_idx]
    for p in cell.paragraphs:
        replace_full_para(p, new_text)
        break


# ─── Main generator ───────────────────────────────────────────────────────────

def generate(data: dict) -> bytes:
    raw = data.get('raw') or {}
    nama = data.get('nama_perusahaan') or raw.get('nama_perusahaan') or NA
    jenis = data.get('jenis_entitas') or raw.get('jenis_entitas') or ''
    periode = data.get('periode') or raw.get('periode') or NA
    hasil_pengawasan = data.get('hasil_pengawasan') or []
    kesimpulan = data.get('kesimpulan') or ''
    tindak_lanjut = data.get('tindak_lanjut') or ''
    nomor_lhptl = raw.get('nomor_lhptl') or NA

    # Derived values
    tahun_ini = periode.strip()
    try:
        tahun_lalu = str(int(tahun_ini) - 1)
    except Exception:
        tahun_lalu = NA

    rekanan_str = terbilang_total(
        raw.get('jumlah_rekanan_perorangan'),
        raw.get('jumlah_rekanan_badan_hukum'),
    )
    pegawai_val = raw.get('jumlah_pegawai')
    pegawai_str = f"{pegawai_val}, termasuk pengurus dan pegawai" if pegawai_val else NA

    afiliasi_list = raw.get('pihak_afiliasi') or []
    afiliasi_str = ', '.join(afiliasi_list) if afiliasi_list else 'Tidak ada'

    tgl_str = date.today().strftime('%-d %B %Y').replace(
        'January', 'Januari').replace('February', 'Februari').replace(
        'March', 'Maret').replace('April', 'April').replace(
        'May', 'Mei').replace('June', 'Juni').replace(
        'July', 'Juli').replace('August', 'Agustus').replace(
        'September', 'September').replace('October', 'Oktober').replace(
        'November', 'November').replace('December', 'Desember')

    # Periode kalimat
    if tahun_ini.isdigit():
        periode_kalimat = f"Periode 1 Januari {tahun_ini} s.d 31 Desember {tahun_ini}."
    else:
        periode_kalimat = f"Periode pengawasan: {tahun_ini}."

    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs

    # ── Title block ──────────────────────────────────────────────────────────
    # Para 1: company name
    replace_full_para(paras[1], nama)
    # Para 2: nomor LHPTL
    replace_full_para(paras[2], f'NOMOR {nomor_lhptl}')

    # ── KV fields (find by label prefix) ─────────────────────────────────────
    kv_map = {
        'nama perusahaan':             nama,
        'alamat':                      raw.get('alamat') or NA,
        'izin usaha':                  raw.get('izin_usaha') or NA,
        'kantor di luar kp':           raw.get('kantor_cabang') or 'Tidak ada',
        'perusahaan afiliasi':         afiliasi_str,
        'jumlah pegawai':              pegawai_str,
        'jumlah rekanan pb':           rekanan_str,
        'kantor akuntan publik':       raw.get('kap_nama') or NA,
        'akuntan publik':              raw.get('akuntan_publik_nama') or NA,
        'nomor izin akuntan':          raw.get('nomor_izin_akuntan') or NA,
        'nomor registrasi akuntan':    raw.get('nomor_registrasi_akuntan') or NA,
        'periode laporan':             tahun_ini,
        'opini':                       raw.get('opini_audit') or NA,
    }

    for para in paras:
        txt = para_full_text(para).strip().lower()
        for label, value in kv_map.items():
            if txt.startswith(label):
                set_last_wt_value(para, value)
                break
        # "Tenaga Kerja Asing: ..." — colon without tab
        if txt.startswith('tenaga kerja asing'):
            tka = raw.get('tka') or 'Tidak ada'
            replace_full_para(para, f'Tenaga Kerja Asing: {tka}')
        # "Penyampaian Laporan Keuangan..." — year reference
        if txt.startswith('penyampaian laporan keuangan') and ('tahun' in txt or 'periode' in txt):
            replace_full_para(para,
                f'Penyampaian Laporan Keuangan yang telah diaudit oleh Akuntan Publik Tahun {tahun_ini}')
        # Laporan Posisi Keuangan year
        if txt.startswith('laporan posisi keuangan') and 'tahun' in txt:
            replace_full_para(para,
                f'Laporan Posisi Keuangan dan Laporan Laba Rugi Komprehensif Tahun {tahun_ini}'
                f' dan {tahun_lalu} sebagai berikut:')
        # Rasio keuangan year
        if txt.startswith('rasio keuangan perusahaan') and 'tahun' in txt:
            replace_full_para(para,
                f'Rasio keuangan Perusahaan untuk periode Tahun {tahun_ini}'
                f' dan {tahun_lalu} adalah sebagai berikut:')
        # Periode pengawasan
        if 'periode 1 januari' in txt or ('periode pengawasan' in txt and 's.d' in txt):
            replace_full_para(para, periode_kalimat)
        # {{PERIODE_KALIMAT}} placeholder jika ada di template
        if '{{periode_kalimat}}' in txt:
            replace_full_para(para, periode_kalimat)
        # Kesimpulan body
        if txt == '...' or txt == '{{kesimpulan}}':
            replace_full_para(para, kesimpulan or NA)
        # Tanda tangan date
        if txt.startswith('jakarta,') and ('{{' in para_full_text(para) or any(y in para_full_text(para) for y in ['2024','2025','2026'])):
            replace_full_para(para, f'Jakarta, {tgl_str}')

    # ── Tables ────────────────────────────────────────────────────────────────
    tables = doc.tables  # noqa — iterable

    # Table 0: Pemegang Saham
    tbl_ps = tables[0]
    clear_data_rows(tbl_ps, keep=1)
    for ps in (raw.get('pemegang_saham') or []):
        add_data_row(tbl_ps, [
            ps.get('nama', '-'),
            fmt_rupiah(ps.get('nilai_rp')),
            f"{ps.get('persentase', 0):.2f}" if ps.get('persentase') is not None else '-',
        ])
    if not (raw.get('pemegang_saham') or []):
        add_data_row(tbl_ps, [NA, '-', '-'])

    # Table 1: Direksi & Komisaris
    tbl_dk = tables[1]
    clear_data_rows(tbl_dk, keep=1)
    for dk in (raw.get('direksi_komisaris') or []):
        add_data_row(tbl_dk, [
            dk.get('jabatan', '-'),
            dk.get('nama', '-'),
            dk.get('surat_persetujuan_ojk', '-') or '-',
        ])
    if not (raw.get('direksi_komisaris') or []):
        add_data_row(tbl_dk, [NA, NA, NA])

    # Table 2: Tenaga Ahli & Pialang
    tbl_ta = tables[2]
    clear_data_rows(tbl_ta, keep=1)
    for ta in (raw.get('tenaga_ahli_pialang') or []):
        ojk_info = ta.get('surat_pengadministrasian_ojk', '') or ''
        reg_no = ta.get('nomor_registrasi', '') or ''
        ojk_combined = f"{reg_no} ({ojk_info})" if ojk_info else reg_no or '-'
        add_data_row(tbl_ta, [
            ta.get('jabatan', '-'),
            ta.get('nama', '-'),
            ojk_combined,
        ])
    if not (raw.get('tenaga_ahli_pialang') or []):
        add_data_row(tbl_ta, [NA, NA, NA])

    # Table 3: Polis Indemnitas
    tbl_pol = tables[3]
    clear_data_rows(tbl_pol, keep=1)
    for pol in (raw.get('polis_indemnitas') or []):
        add_data_row(tbl_pol, [
            pol.get('nomor_polis', '-'),
            pol.get('penanggung', '-'),
            fmt_rupiah(pol.get('nilai_pertanggungan')),
            pol.get('masa_berlaku', '-'),
        ])
    if not (raw.get('polis_indemnitas') or []):
        add_data_row(tbl_pol, ['-', '-', '-', '-'])

    # Table 4: Sanksi
    tbl_sank = tables[4]
    clear_data_rows(tbl_sank, keep=1)
    for s in (raw.get('sanksi') or []):
        add_data_row(tbl_sank, [
            '-',
            s.get('nomor_tanggal', '-'),
            s.get('jenis', '-'),
            s.get('penyebab', '-'),
        ])
    if not (raw.get('sanksi') or []):
        add_data_row(tbl_sank, ['-', '-', '-', '-'])

    # Table 5: Neraca + Laba Rugi (2 header rows; update year in header)
    tbl_lk = tables[5]
    update_header_cell(tbl_lk, 0, 1, f'Tahun {tahun_ini}')
    update_header_cell(tbl_lk, 0, 2, f'Tahun {tahun_lalu}')
    clear_data_rows(tbl_lk, keep=2)
    for row in (raw.get('neraca_laba_rugi') or []):
        add_data_row(tbl_lk, [
            row.get('label', '-'),
            fmt_rupiah(row.get('nilai_ini')),
            fmt_rupiah(row.get('nilai_lalu')),
        ], header_row_idx=0)
    if not (raw.get('neraca_laba_rugi') or []):
        add_data_row(tbl_lk, ['Data tidak tersedia', '-', '-'], header_row_idx=0)

    # Table 6: Rasio Keuangan
    tbl_rasio = tables[6]
    update_header_cell(tbl_rasio, 0, 1, f'Tahun {tahun_ini}')
    update_header_cell(tbl_rasio, 0, 2, f'Tahun {tahun_lalu}')
    clear_data_rows(tbl_rasio, keep=1)
    for row in (raw.get('rasio_keuangan_tabel') or []):
        add_data_row(tbl_rasio, [
            row.get('label', '-'),
            fmt_pct(row.get('nilai_ini')),
            fmt_pct(row.get('nilai_lalu')),
        ])
    if not (raw.get('rasio_keuangan_tabel') or []):
        add_data_row(tbl_rasio, ['Data tidak tersedia', '-', '-'])

    # Table 7: Tim Pengawas — biarkan kosong (diisi manual)
    # (tidak diubah)

    # Table 8: Hasil Pengawasan
    tbl_hasil = tables[8]
    clear_data_rows(tbl_hasil, keep=1)
    tipe_label = {'pelanggaran': 'Pelanggaran', 'perlu_perhatian': 'Perlu Perhatian'}
    for h in hasil_pengawasan:
        add_data_row(tbl_hasil, [
            str(h.get('nomor', '')),
            h.get('catatan', '-'),
            tipe_label.get(h.get('tipe', ''), ''),
        ])
    if not hasil_pengawasan:
        add_data_row(tbl_hasil, ['-', 'Tidak ada temuan', ''])

    # Table 9: Tindak Lanjut
    tbl_tl = tables[9]
    clear_data_rows(tbl_tl, keep=1)
    # Parse tindak_lanjut: "1. Teks\n2. Teks" → poin bersih
    import re as _re
    poin_tl = [p.strip() for p in _re.split(r'(?m)^\d+\.\s*', tindak_lanjut) if p.strip()]
    # Fallback: split by newline
    if not poin_tl:
        poin_tl = [p.strip() for p in tindak_lanjut.split('\n') if p.strip()]
    for i, poin in enumerate(poin_tl, 1):
        add_data_row(tbl_tl, [f'{i}.', poin, ''])
    if not poin_tl:
        add_data_row(tbl_tl, ['1.', NA, ''])

    # Table 10: TTD — biarkan (tidak diubah)

    # ── Output ────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == '__main__':
    if len(sys.argv) > 1:
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        data = json.load(sys.stdin)

    result = generate(data)
    sys.stdout.buffer.write(result)
